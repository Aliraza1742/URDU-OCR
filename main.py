import itertools
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert
from pdf2image import convert_from_path
from PIL import Image
import csv
import os

# Urdu alphabet (basic)
urdu_alphabet = ['ا', 'ب', 'پ', 'ت', 'ٹ', 'ث', 'ج', 'چ', 'ح', 'خ', 'د', 'ڈ', 'ذ', 'ر', 'ڑ', 'ز', 'ژ', 'س', 'ش', 'ص', 'ض', 'ط', 'ظ', 'ع', 'غ', 'ف', 'ق', 'ک', 'گ', 'ل', 'م', 'ن', 'و', 'ہ', 'ء', 'ی']

def generate_urdu_words():
    start_letter = 'ط'
    middle_letters = ['ا', 'ل', 'م']
    end_letters = ['ہ', 'ن', 'ی']
    words = []

    for mid1, mid2, end in itertools.product(middle_letters, repeat=2, r=end_letters):
        word = start_letter + mid1 + mid2 + end
        index = [urdu_alphabet.index(c) for c in word if c in urdu_alphabet]
        label = '_'.join(map(str, index))
        words.append((word, label))

    return words

def create_word_document(words, filename="Urdu_Words_Dataset.docx"):
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    table = doc.add_table(rows=0, cols=4)
    font_name = 'Noori Nastaleeq'  # Ensure it's installed

    for i, (word, _) in enumerate(words):
        if i % 4 == 0:
            row = table.add_row().cells

        cell = row[i % 4]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(word)
        run.font.size = Pt(28)
        run.font.name = font_name
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.save(filename)

def generate_labels_csv(indexed_words, filename="Urdu_Words_Labels.csv"):
    with open(filename, mode='w', newline='', encoding='utf-8') as file:
        writer = csv.writer(file)
        writer.writerow(["Label"])
        for _, label in indexed_words:
            writer.writerow([label])

def process_images_folder(images_folder="png_images", output_folder="output", labels_csv="Urdu_Words_Labels.csv"):
    os.makedirs(output_folder, exist_ok=True)
    with open(labels_csv, newline='', encoding='utf-8') as file:
        reader = csv.DictReader(file)
        labels = [row["Label"] for row in reader]

    word_index = 0
    for filename in sorted(os.listdir(images_folder)):
        if filename.endswith(".png"):
            image_path = os.path.join(images_folder, filename)
            image = Image.open(image_path)
            img_width, img_height = image.size

            word_width = 300
            line_height = 150
            margin_left = 100
            margin_top = 50
            spaces_between_words = 75

            words_per_line = 4
            lines = img_height // line_height

            for i in range(lines):
                for j in range(words_per_line):
                    if word_index >= len(labels):
                        return
                    left = margin_left + j * (word_width + spaces_between_words)
                    upper = margin_top + i * line_height
                    right = left + word_width
                    lower = upper + line_height

                    word_img = image.crop((left, upper, right, lower))
                    label = labels[word_index]
                    word_img.save(os.path.join(output_folder, f"{label}.png"))
                    word_index += 1

# ---------- Pipeline Execution ----------

# Step 1: Generate words and labels
indexed_words = generate_urdu_words()

# Step 2: Create DOCX
create_word_document(indexed_words, "Urdu_Words_Dataset.docx")

# Step 3: Convert DOCX to PDF
convert("Urdu_Words_Dataset.docx", "Urdu_Words_Dataset.pdf")

# Step 4: Convert PDF to PNG
os.makedirs("png_images", exist_ok=True)
images = convert_from_path("Urdu_Words_Dataset.pdf", output_folder="png_images")

# Step 5: Generate CSV labels
generate_labels_csv(indexed_words, "Urdu_Words_Labels.csv")

# Step 6: Crop and save word images
process_images_folder(images_folder="png_images", output_folder="output", labels_csv="Urdu_Words_Labels.csv")
